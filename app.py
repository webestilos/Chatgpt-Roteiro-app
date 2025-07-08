import streamlit as st
import pyautogui
import pyperclip
import time
import webbrowser
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Roteiro ChatGPT", layout="centered")
st.title("🎬 Criador de Roteiro com ChatGPT")

if "historico" not in st.session_state:
    st.session_state.historico = []

# Campo visível, mas não será salvo no DOC
roteirista = st.text_area("🖊️ Roteirista (nome, canal ou assinatura):", height=100)

# Campo apresentação + tema com texto orientador
texto_padrao_apresentacao = """2 – Apresentação leve e reflexiva do tema, após a apresentação do tema faça as pessoas comentarem.

Hoje quero falar sobre ..., o tema:"""

apresentacao = st.text_area(
    "🎙️ Apresentação + Tema:",
    value=texto_padrao_apresentacao,
    height=150
)

# Campo tópicos com instrução pré-preenchida
texto_padrao_topicos = """3 – Exploração dos princípios com profundidade e calma.
Um a um, em tom de meditação e despertar. Escreva o primeiro e peça para prosseguir:"""

topicos = st.text_area(
    "📝 Tópicos do vídeo (um por linha):",
    value=texto_padrao_topicos,
    height=200
)

# Campo de conclusão
conclusao = st.text_area("🏁 5 – Conclusão com convite para viver com consciência e gratidão:", height=80)

# Monta o prompt para o ChatGPT
def montar_prompt(roteirista, apresentacao, topicos, conclusao):
    prompt = (
        f"Roteirista: {roteirista}\n\n"
        f"{apresentacao}\n\n"
        f"{topicos}\n\n"
        f"{conclusao}"
    )
    return prompt

# Enviar para o ChatGPT via automação
if st.button("🚀 Enviar para ChatGPT"):
    if any(campo.strip() == "" for campo in [roteirista, apresentacao, topicos, conclusao]):
        st.warning("Preencha todos os campos antes de enviar.")
    else:
        prompt = montar_prompt(roteirista, apresentacao, topicos, conclusao)
        pyperclip.copy(prompt)
        webbrowser.open("https://chat.openai.com/")
        time.sleep(8)
        pyautogui.hotkey("ctrl", "v")
        time.sleep(0.5)
        pyautogui.press("enter")
        st.success("Prompt enviado com sucesso! Cole a resposta abaixo para salvar o roteiro.")

# Resposta gerada pelo ChatGPT
resposta = st.text_area("📋 Cole a resposta do ChatGPT aqui:", height=300)

# Salvar o roteiro completo, sem o campo "roteirista"
if st.button("💾 Salvar roteiro e gerar arquivo para download"):
    if resposta.strip() == "":
        st.warning("Cole a resposta do ChatGPT para salvar o roteiro.")
    else:
        st.session_state.historico.append({
            "apresentacao": apresentacao,
            "topicos": topicos,
            "conclusao": conclusao,
            "resposta": resposta
        })

        # Criar o DOC sem incluir o roteirista
        doc = Document()
        doc.add_heading("🎬 Roteiro de Vídeo", 0)
        doc.add_paragraph(" ")

        for i, item in enumerate(st.session_state.historico, 1):
            doc.add_heading(f"{i}. Roteiro Completo", level=1)
            doc.add_paragraph(item["resposta"])
            doc.add_paragraph("\n")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("Arquivo .docx criado com sucesso!")

        st.download_button(
            label="⬇️ Baixar roteiro",
            data=buffer,
            file_name="roteiro_chatgpt.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
